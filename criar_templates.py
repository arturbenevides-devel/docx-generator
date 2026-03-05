"""
Script para criar os templates .docx com placeholders do docxtpl.
Execute uma vez: python3 criar_templates.py
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def add_paragraph(doc, text, bold=False, alignment=None, font_size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    if alignment:
        p.alignment = alignment
    return p


def criar_notificacao_sem_fiador():
    doc = Document()

    add_paragraph(doc, "São Paulo, {{ hoje_extenso }}")
    add_paragraph(doc, "")
    add_paragraph(doc, "Ilmo(a). Sr.(a)")
    add_paragraph(doc, "")

    add_paragraph(doc, "DEVEDOR(A):", bold=True)
    add_paragraph(doc, "Nome: {{ nome_pes }} - CPF: {{ cpf_formatado }}")
    add_paragraph(doc, "{{ Endereco_pend }}, nº {{ NumEnd_pend }}{% if ComplEndereco_pend %} - {{ ComplEndereco_pend }}{% endif %}")
    add_paragraph(doc, "Bairro {{ Bairro_pend }} – {{ Cidade_pend }} – {{ UF_pend }} – CEP: {{ CEP_pend }}")
    add_paragraph(doc, "")

    add_paragraph(doc,
        'A (razao social credor), com sede social na (endereço credor) - (bairro credor) - '
        '(cidade credor) - (estado credor) - (cep credor) vem através da presente apresentar '
        'NOTIFICAÇÃO EXTRAJUDICIAL, pelas razões a seguir aduzidas:'
    )
    add_paragraph(doc, "")

    add_paragraph(doc,
        'O NOTIFICADO assinou em (data assinatura) Instrumento Particular de Compromisso de '
        'Venda e Compra de Unidade Autônoma referente ao empreendimento {{ obra_nome }}, '
        'localizado à {{ obra_endereco }}, bairro {{ obra_bairro }}, na cidade de '
        '{{ obra_cidade }} - {{ obra_uf }}, CEP {{ obra_cep }}, referente à unidade {{ ConcatIdentificador_unid }}.'
    )
    add_paragraph(doc, "")

    add_paragraph(doc,
        'No citado instrumento, dispõe que os COMPRADORES constituído em mora, e não efetuada '
        'por ele a purga integral e com todos os encargos contratuais e legalmente previstos, '
        'no prazo de 15 (quinze) dias, terá a "VENDEDORA" o direito de adotar as providencias, '
        'nos termos estabelecidos no referido instrumento.'
    )
    add_paragraph(doc, "")

    add_paragraph(doc,
        'Isto posto, informamos que Vossa Senhoria possui parcelas vencidas, portanto, pela '
        'presente, NOTIFICAMOS V.S. para que regularize as pendências necessárias para manutenção '
        'do contrato no prazo improrrogável de 15 (quinze) dias, sob pena de estar constituído '
        'em mora e incidir em infração contratual e nas penalidades nele previstas para a mora '
        'quanto às obrigações contraídas. O não atendimento da presente notificação no prazo '
        'concedido importará na adoção das medidas legais cabíveis, independentemente de nova notificação.'
    )
    add_paragraph(doc, "")

    add_paragraph(doc, "Certos de que seremos prontamente atendidos, era o que nos cumpria")
    add_paragraph(doc, "")
    add_paragraph(doc, "Atenciosamente,")

    return doc


def criar_notificacao_com_fiador():
    doc = Document()

    add_paragraph(doc, "São Paulo, {{ hoje_extenso }}")
    add_paragraph(doc, "")
    add_paragraph(doc, "Ilmo(a). Sr.(a)")
    add_paragraph(doc, "")

    add_paragraph(doc, "DEVEDOR(A):", bold=True)
    add_paragraph(doc, "Nome: {{ nome_pes }} - CPF: {{ cpf_formatado }}")
    add_paragraph(doc, "{{ Endereco_pend }}, nº {{ NumEnd_pend }}{% if ComplEndereco_pend %} - {{ ComplEndereco_pend }}{% endif %}")
    add_paragraph(doc, "Bairro {{ Bairro_pend }} – {{ Cidade_pend }} – {{ UF_pend }} – CEP: {{ CEP_pend }}")
    add_paragraph(doc, "")

    # Bloco do fiador com loop do docxtpl
    add_paragraph(doc, "AVALISTA:", bold=True)
    add_paragraph(doc, "Nome: {{ Fiador }} - CPF: {{ fiador_cpf }}")
    add_paragraph(doc, "{{ fiador_endereco }}, nº {{ fiador_num_endereco }}{% if fiador_complemento %} - {{ fiador_complemento }}{% endif %}")
    add_paragraph(doc, "Bairro {{ fiador_bairro }} – {{ fiador_cidade }} – {{ fiador_uf }} – CEP: {{ fiador_cep }}")
    add_paragraph(doc, "")

    add_paragraph(doc,
        'A (razao social credor), com sede social na (endereço credor) - (bairro credor) - '
        '(cidade credor) - (estado credor) - (cep credor) vem através da presente apresentar '
        'NOTIFICAÇÃO EXTRAJUDICIAL, pelas razões a seguir aduzidas:'
    )
    add_paragraph(doc, "")

    add_paragraph(doc,
        'O NOTIFICADO assinou em (data assinatura) Instrumento Particular de Compromisso de '
        'Venda e Compra de Unidade Autônoma referente ao empreendimento {{ obra_nome }}, '
        'localizado à {{ obra_endereco }}, bairro {{ obra_bairro }}, na cidade de '
        '{{ obra_cidade }} - {{ obra_uf }}, CEP {{ obra_cep }}, referente à unidade {{ ConcatIdentificador_unid }}.'
    )
    add_paragraph(doc, "")

    add_paragraph(doc,
        'No citado instrumento, dispõe que os COMPRADORES constituído em mora, e não efetuada '
        'por ele a purga integral e com todos os encargos contratuais e legalmente previstos, '
        'no prazo de 15 (quinze) dias, terá a "VENDEDORA" o direito de adotar as providencias, '
        'nos termos estabelecidos no referido instrumento.'
    )
    add_paragraph(doc, "")

    add_paragraph(doc,
        'Isto posto, informamos que Vossa Senhoria possui parcelas vencidas, portanto, pela '
        'presente, NOTIFICAMOS V.S. para que regularize as pendências necessárias para manutenção '
        'do contrato no prazo improrrogável de 15 (quinze) dias, sob pena de estar constituído '
        'em mora e incidir em infração contratual e nas penalidades nele previstas para a mora '
        'quanto às obrigações contraídas. O não atendimento da presente notificação no prazo '
        'concedido importará na adoção das medidas legais cabíveis, independentemente de nova notificação.'
    )
    add_paragraph(doc, "")

    add_paragraph(doc, "Certos de que seremos prontamente atendidos, era o que nos cumpria")
    add_paragraph(doc, "")
    add_paragraph(doc, "Atenciosamente,")

    return doc


def criar_execucao_sem_fiador():
    doc = Document()

    add_paragraph(doc,
        'EXCELENTÍSSIMO SENHOR DOUTOR JUIZ DE DIREITO DA M.M. VARA CÍVEL DO FORO DE ................ - SP',
        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER
    )
    add_paragraph(doc, "")

    add_paragraph(doc,
        '........................................, pessoa jurídica de direito privado devidamente inscrita no '
        'CNPJ sob o nº ............../........-....., com escritório administrativo na (rua credor), nº (numero credor), '
        '(complemento credor), Bairro (bairro credor) – (cep credor), endereço de e-mail: (email credor), '
        'por seus advogados e procuradores, abaixo assinados, vem mui respeitosamente, propor a presente ação de '
        'execução por quantia certa em face de {{ nome_pes }}, (nacionalidade), (estado civil), (profissão), '
        'portador da cédula de identidade nº ................ e inscrito no CPF sob o nº {{ cpf_formatado }}, '
        'residente e domiciliado na {{ Endereco_pend }}, nº {{ NumEnd_pend }} – Bairro {{ Bairro_pend }} – '
        'Cidade {{ Cidade_pend }} – {{ UF_pend }}, CEP {{ CEP_pend }}, '
        'pelas razões de fato e de direito a seguir expostas:'
    )
    add_paragraph(doc, "")

    # Itens 1 a 12
    add_paragraph(doc,
        '1) – A Exequente se constitui em sociedade de propósitos específicos (SPE), cujo objeto é a incorporação '
        'imobiliária e construção de empreendimento denominado {{ obra_nome }}, no município de {{ obra_cidade }}'
    )

    add_paragraph(doc,
        '2) – Dando cumprimento ao seu objeto, incorporou e construiu edifício de apartamentos residenciais, '
        'onde os Executados adquiriram a Unidade {{ ConcatIdentificador_unid }} Parte do preço do imóvel foi pago '
        'através de financiamento bancário. O saldo remanescente, não coberto pelo valor financiado, foi alvo de '
        'negociação entre as partes para pagamento parcelado.'
    )

    add_paragraph(doc,
        'Referida negociação foi inserida em instrumento de confissão de dívida firmada entre as partes em '
        '.. de ...... de ...., onde os executados confessam dever para a Exequente o valor de R$ ...... (.......) '
        'e se compromete a quitar referida dívida na forma descrita no respectivo instrumento.'
    )

    add_paragraph(doc,
        '3) – Ocorre que os Devedores não cumpriram com as obrigações pactuadas, deixando de efetuar o pagamento '
        'das parcelas mensais/anuais a partir de ../../...., totalizando até ../../...., atualizado, R$ ...... (.......).'
    )

    add_paragraph(doc,
        'Para casos de inadimplemento, o contrato prevê aplicação de multa de 2% sobre o valor da dívida, '
        'além de correção monetária e juros de mora de 1% ao mês. O instrumento prevê, ainda, que em caso de '
        'inadimplemento de ..................., consecutivas ou não, haverá o vencimento antecipado de todas '
        'as parcelas vincendas, podendo a dívida ser integralmente cobrada em uma única parcela.'
    )

    add_paragraph(doc,
        '4) – Além dos valores das parcelas do saldo do preço do imóvel, os Executados deixaram de efetuar '
        'o pagamento da denominada "Parcela da Evolução de Obras", previsto no contrato de compra e venda e '
        'financiamento de construção, no qual a Exequente figura como fiadora dos Executados, se responsabilizando '
        'pelo adimplemento das ditas parcelas. Ao longo do contrato, os executados deixaram de quitar as parcelas '
        'descritas no extrato financeiro, que segue o presente.'
    )

    add_paragraph(doc,
        'Desta feita, a Autora suportou o pagamento da importância nominal de R$ ...... (.......), valor este '
        'sem correção, mediante descontos sofridos nas medições pagas pela CEF. Nos termos do contrato firmado, '
        'a Incorporadora/Construtora, ora Autora tem o direito de cobrar a devedora do contrato de financiamento.'
    )

    add_paragraph(doc,
        'Nos termos da planilha que instrui a presente, o débito atualizado, até a data de ../../...., '
        'totaliza o montante de R$ ...... (.........).'
    )

    add_paragraph(doc,
        '5) – Ocorre que a Ré deixou de efetuar o pagamento da parcela de correção monetária pelo INCC '
        '(Correção Idx), nos exatos termos do pacto havido em contrato, tornando-se devedora da quantia de '
        'R$ ...... (.......), valor este sendo o valor principal, conforme se demonstra na planilha de apuração '
        'anexada ao presente. O valor atualizado até a data de ../../.... é de R$ ...... (.......).'
    )

    add_paragraph(doc,
        '6) – Além disso, os Devedores deixaram também de realizar o pagamento referente ao custo de '
        'atribuição de unidade, que corresponde, até ../../...., R$ ...... (.......)'
    )

    add_paragraph(doc,
        'É necessário esclarecimento essenciais para a compreensão da cláusula contratual que trata das '
        'despesas cartorárias com registros e averbações. Conforme regra do artigo 490 do Código Civil, o '
        'comprador é responsável pelo pagamento das despesas com registro de imóvel adquirido.'
    )

    add_paragraph(doc,
        '7) – Além disso, o devedor deixou de efetuar o pagamento da taxa condominial, cujo valor, '
        'até ../../...., é de R$ ........ (........................).'
    )

    add_paragraph(doc,
        'É necessário o esclarecimento de pontos essenciais para a compreensão da parcela devida, a fim de '
        'demonstrar que a taxa é devida, considerando que o réu recebeu o imóvel em ..... de ........... de .........'
    )

    add_paragraph(doc,
        '8) – Os devedores também não efetuaram o reembolso das despesas de registro imobiliário '
        'no valor de R$........ (......) O valor atualizado até ......., conforme planilha, é de R$ ........ (..........).'
    )

    add_paragraph(doc,
        '9) – Diante do inadimplemento parcial da dívida confessada, consubstanciado no atraso das parcelas '
        'mensais e/ou anuais, bem como das parcelas referentes à Evolução de Obra, INCC, custo de atribuição '
        'de matrícula, taxas condominiais e despesas de registro imobiliário, considera-se integralmente vencido '
        'o saldo devedor, passível de cobrança pela via executiva.'
    )

    add_paragraph(doc,
        '10) – O artigo 784 do Código de Processo Civil traz rol dos títulos executivos extrajudiciais, '
        'donde destacamos documento particular assinado pelo devedor e por duas testemunhas (inc. III).'
    )

    add_paragraph(doc,
        'Portanto, a dívida confessada e não liquidada pode ser objeto de ação executiva, tratando-se de '
        'título executivo extrajudicial, líquida, certa e exigível.'
    )

    add_paragraph(doc,
        '11) – O valor total da presente execução, já abatendo as parcelas pagas pelos Executados, '
        'é de R$ ............ (..........), conforme planilha de cálculos que segue anexada ao presente.'
    )

    add_paragraph(doc,
        '12) – A Exequente buscou, insistentemente, negociar com os Devedores o recebimento da dívida '
        'pelos meios extrajudiciais, sem que houvesse sucesso. Diante do inadimplemento, a Credora se vale '
        'do presente processo executivo para receber seu crédito.'
    )
    add_paragraph(doc, "")

    add_paragraph(doc, "DO PEDIDO", bold=True)
    add_paragraph(doc, "")

    add_paragraph(doc,
        'Diante do exposto, requer a V.Exa. a citação dos Executados através de carta postal para os termos '
        'da presente execução, para que paguem o débito apurado de R$ ............ (..........), acrescidos das '
        'custas processuais e honorários de advogado, no prazo de 3 (três) dias, sob pena de lhe serem '
        'penhorados tantos bens quanto bastem para solução do débito. Sobre referido valor deverá incidir '
        'correção monetária até a data do efetivo pagamento e juros de mora de 1% ao mês a partir da citação.'
    )

    add_paragraph(doc,
        'Se os Devedores não efetuarem o pagamento, a Exequente indica, desde logo, que a penhora recaia '
        'em dinheiro, através do sistema Bacenjud. Requer que a intimação da penhora ocorra através de '
        'advogado constituído nos autos pelos Executados (CPC, art. 841) ou por via postal, caso não tenha '
        'advogado constituído.'
    )

    add_paragraph(doc,
        'Requer, outrossim, a fixação de honorários para a presente execução, nos termos do art. 827 do '
        'Código de Processo Civil, que serão reduzidos pela metade no caso de pagamento em 3 (três) dias da citação.'
    )

    add_paragraph(doc,
        'Por derradeiro, requer-se que todas as publicações e intimações do presente feito em nome de '
        'Claudio Oliveira Cabral Jr, advogado, regularmente inscrito na OAB/SP nº 130-544 e '
        'Gislene Caetano de Oliveira Andres, advogada regularmente inscrita na OAB/SP nº 192-104, '
        'com escritório profissional na Praça da Sé, nº 399, sala 02, Bairro Sé – CEP: 01001-000'
    )

    add_paragraph(doc, 'Dá-se à causa o valor de R$ ............ (..........)')
    add_paragraph(doc, "")

    add_paragraph(doc, "Termos em que,")
    add_paragraph(doc, "P. Deferimento.")
    add_paragraph(doc, "")
    add_paragraph(doc, "São Paulo, {{ hoje_extenso }}")
    add_paragraph(doc, "")
    add_paragraph(doc, "Claudio Oliveira Cabral Jr.")
    add_paragraph(doc, "OAB/SP 130.544")
    add_paragraph(doc, "")
    add_paragraph(doc, "Gislene Caetano de Oliveira Andres")
    add_paragraph(doc, "OAB/SP 192.104")

    return doc


def criar_execucao_com_fiador():
    doc = Document()

    add_paragraph(doc,
        'EXCELENTÍSSIMO SENHOR DOUTOR JUIZ DE DIREITO DA M.M. VARA CÍVEL DO FORO DE ................ - SP',
        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER
    )
    add_paragraph(doc, "")

    add_paragraph(doc,
        '........................................, pessoa jurídica de direito privado devidamente inscrita no '
        'CNPJ sob o nº ............../........-....., com escritório administrativo na (rua credor), nº (numero credor), '
        '(complemento credor), Bairro (bairro credor) – (cep credor), endereço de e-mail: (email credor), '
        'por seus advogados e procuradores, abaixo assinados, vem mui respeitosamente, propor a presente ação de '
        'execução por quantia certa em face de {{ nome_pes }}, (nacionalidade), (estado civil), (profissão), '
        'portador da cédula de identidade nº ................ e inscrito no CPF sob o nº {{ cpf_formatado }}, '
        'residente e domiciliado na {{ Endereco_pend }}, nº {{ NumEnd_pend }} – Bairro {{ Bairro_pend }} – '
        'Cidade {{ Cidade_pend }} – {{ UF_pend }}, CEP {{ CEP_pend }}, '
        'e {{ Fiador }}, inscrito no CPF sob o nº {{ fiador_cpf }}, '
        'residente e domiciliado na {{ fiador_endereco }}, nº {{ fiador_num_endereco }} – '
        'Bairro {{ fiador_bairro }} – {{ fiador_uf }}, CEP {{ fiador_cep }}, '
        'na qualidade de FIADOR(A)/AVALISTA, '
        'pelas razões de fato e de direito a seguir expostas:'
    )
    add_paragraph(doc, "")

    # Itens 1 a 12 (mesmo conteúdo do sem_fiador)
    add_paragraph(doc,
        '1) – A Exequente se constitui em sociedade de propósitos específicos (SPE), cujo objeto é a incorporação '
        'imobiliária e construção de empreendimento denominado {{ obra_nome }}, no município de {{ obra_cidade }}'
    )

    add_paragraph(doc,
        '2) – Dando cumprimento ao seu objeto, incorporou e construiu edifício de apartamentos residenciais, '
        'onde os Executados adquiriram a Unidade {{ ConcatIdentificador_unid }} Parte do preço do imóvel foi pago '
        'através de financiamento bancário. O saldo remanescente, não coberto pelo valor financiado, foi alvo de '
        'negociação entre as partes para pagamento parcelado.'
    )

    add_paragraph(doc,
        'Referida negociação foi inserida em instrumento de confissão de dívida firmada entre as partes em '
        '.. de ...... de ...., onde os executados confessam dever para a Exequente o valor de R$ ...... (.......) '
        'e se compromete a quitar referida dívida na forma descrita no respectivo instrumento.'
    )

    add_paragraph(doc,
        '3) – Ocorre que os Devedores não cumpriram com as obrigações pactuadas, deixando de efetuar o pagamento '
        'das parcelas mensais/anuais a partir de ../../...., totalizando até ../../...., atualizado, R$ ...... (.......).'
    )

    add_paragraph(doc,
        'Para casos de inadimplemento, o contrato prevê aplicação de multa de 2% sobre o valor da dívida, '
        'além de correção monetária e juros de mora de 1% ao mês. O instrumento prevê, ainda, que em caso de '
        'inadimplemento de ..................., consecutivas ou não, haverá o vencimento antecipado de todas '
        'as parcelas vincendas, podendo a dívida ser integralmente cobrada em uma única parcela.'
    )

    add_paragraph(doc,
        '4) – Além dos valores das parcelas do saldo do preço do imóvel, os Executados deixaram de efetuar '
        'o pagamento da denominada "Parcela da Evolução de Obras", previsto no contrato de compra e venda e '
        'financiamento de construção, no qual a Exequente figura como fiadora dos Executados, se responsabilizando '
        'pelo adimplemento das ditas parcelas. Ao longo do contrato, os executados deixaram de quitar as parcelas '
        'descritas no extrato financeiro, que segue o presente.'
    )

    add_paragraph(doc,
        'Desta feita, a Autora suportou o pagamento da importância nominal de R$ ...... (.......), valor este '
        'sem correção, mediante descontos sofridos nas medições pagas pela CEF. Nos termos do contrato firmado, '
        'a Incorporadora/Construtora, ora Autora tem o direito de cobrar a devedora do contrato de financiamento.'
    )

    add_paragraph(doc,
        'Nos termos da planilha que instrui a presente, o débito atualizado, até a data de ../../...., '
        'totaliza o montante de R$ ...... (.........).'
    )

    add_paragraph(doc,
        '5) – Ocorre que a Ré deixou de efetuar o pagamento da parcela de correção monetária pelo INCC '
        '(Correção Idx), nos exatos termos do pacto havido em contrato, tornando-se devedora da quantia de '
        'R$ ...... (.......), valor este sendo o valor principal, conforme se demonstra na planilha de apuração '
        'anexada ao presente. O valor atualizado até a data de ../../.... é de R$ ...... (.......).'
    )

    add_paragraph(doc,
        '6) – Além disso, os Devedores deixaram também de realizar o pagamento referente ao custo de '
        'atribuição de unidade, que corresponde, até ../../...., R$ ...... (.......)'
    )

    add_paragraph(doc,
        'É necessário esclarecimento essenciais para a compreensão da cláusula contratual que trata das '
        'despesas cartorárias com registros e averbações. Conforme regra do artigo 490 do Código Civil, o '
        'comprador é responsável pelo pagamento das despesas com registro de imóvel adquirido.'
    )

    add_paragraph(doc,
        '7) – Além disso, o devedor deixou de efetuar o pagamento da taxa condominial, cujo valor, '
        'até ../../...., é de R$ ........ (........................).'
    )

    add_paragraph(doc,
        'É necessário o esclarecimento de pontos essenciais para a compreensão da parcela devida, a fim de '
        'demonstrar que a taxa é devida, considerando que o réu recebeu o imóvel em ..... de ........... de .........'
    )

    add_paragraph(doc,
        '8) – Os devedores também não efetuaram o reembolso das despesas de registro imobiliário '
        'no valor de R$........ (......) O valor atualizado até ......., conforme planilha, é de R$ ........ (..........).'
    )

    add_paragraph(doc,
        '9) – Diante do inadimplemento parcial da dívida confessada, consubstanciado no atraso das parcelas '
        'mensais e/ou anuais, bem como das parcelas referentes à Evolução de Obra, INCC, custo de atribuição '
        'de matrícula, taxas condominiais e despesas de registro imobiliário, considera-se integralmente vencido '
        'o saldo devedor, passível de cobrança pela via executiva.'
    )

    add_paragraph(doc,
        '10) – O artigo 784 do Código de Processo Civil traz rol dos títulos executivos extrajudiciais, '
        'donde destacamos documento particular assinado pelo devedor e por duas testemunhas (inc. III).'
    )

    add_paragraph(doc,
        'Portanto, a dívida confessada e não liquidada pode ser objeto de ação executiva, tratando-se de '
        'título executivo extrajudicial, líquida, certa e exigível.'
    )

    add_paragraph(doc,
        '11) – O valor total da presente execução, já abatendo as parcelas pagas pelos Executados, '
        'é de R$ ............ (..........), conforme planilha de cálculos que segue anexada ao presente.'
    )

    add_paragraph(doc,
        '12) – A Exequente buscou, insistentemente, negociar com os Devedores o recebimento da dívida '
        'pelos meios extrajudiciais, sem que houvesse sucesso. Diante do inadimplemento, a Credora se vale '
        'do presente processo executivo para receber seu crédito.'
    )
    add_paragraph(doc, "")

    add_paragraph(doc, "DO PEDIDO", bold=True)
    add_paragraph(doc, "")

    add_paragraph(doc,
        'Diante do exposto, requer a V.Exa. a citação dos Executados através de carta postal para os termos '
        'da presente execução, para que paguem o débito apurado de R$ ............ (..........), acrescidos das '
        'custas processuais e honorários de advogado, no prazo de 3 (três) dias, sob pena de lhe serem '
        'penhorados tantos bens quanto bastem para solução do débito. Sobre referido valor deverá incidir '
        'correção monetária até a data do efetivo pagamento e juros de mora de 1% ao mês a partir da citação.'
    )

    add_paragraph(doc,
        'Se os Devedores não efetuarem o pagamento, a Exequente indica, desde logo, que a penhora recaia '
        'em dinheiro, através do sistema Bacenjud. Requer que a intimação da penhora ocorra através de '
        'advogado constituído nos autos pelos Executados (CPC, art. 841) ou por via postal, caso não tenha '
        'advogado constituído.'
    )

    add_paragraph(doc,
        'Requer, outrossim, a fixação de honorários para a presente execução, nos termos do art. 827 do '
        'Código de Processo Civil, que serão reduzidos pela metade no caso de pagamento em 3 (três) dias da citação.'
    )

    add_paragraph(doc,
        'Por derradeiro, requer-se que todas as publicações e intimações do presente feito em nome de '
        'Claudio Oliveira Cabral Jr, advogado, regularmente inscrito na OAB/SP nº 130-544 e '
        'Gislene Caetano de Oliveira Andres, advogada regularmente inscrita na OAB/SP nº 192-104, '
        'com escritório profissional na Praça da Sé, nº 399, sala 02, Bairro Sé – CEP: 01001-000'
    )

    add_paragraph(doc, 'Dá-se à causa o valor de R$ ............ (..........)')
    add_paragraph(doc, "")

    add_paragraph(doc, "Termos em que,")
    add_paragraph(doc, "P. Deferimento.")
    add_paragraph(doc, "")
    add_paragraph(doc, "São Paulo, {{ hoje_extenso }}")
    add_paragraph(doc, "")
    add_paragraph(doc, "Claudio Oliveira Cabral Jr.")
    add_paragraph(doc, "OAB/SP 130.544")
    add_paragraph(doc, "")
    add_paragraph(doc, "Gislene Caetano de Oliveira Andres")
    add_paragraph(doc, "OAB/SP 192.104")

    return doc


if __name__ == "__main__":

    templates = {
        "templates/notificacao/sem_fiador.docx": criar_notificacao_sem_fiador,
        "templates/notificacao/com_fiador.docx": criar_notificacao_com_fiador,
        "templates/execucao/sem_fiador.docx": criar_execucao_sem_fiador,
        "templates/execucao/com_fiador.docx": criar_execucao_com_fiador,
    }

    for caminho, criar_func in templates.items():
        os.makedirs(os.path.dirname(caminho), exist_ok=True)
        doc = criar_func()
        doc.save(caminho)
        print(f"✅ Criado: {caminho}")

    print("\nTodos os templates foram criados com sucesso!")
